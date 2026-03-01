import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

class ReportGenerator:
    """Generates professional assessment reports in DOCX format."""
    
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, client_name, assessment_type, findings):
        doc = Document()
        
        # Title Page
        doc.add_heading('Security Assessment Report', 0)
        doc.add_paragraph(f"Client: {client_name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Assessment Type: {assessment_type}")
        
        doc.add_page_break()

        # Findings Summary
        doc.add_heading('Findings Summary', level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Severity'
        hdr_cells[2].text = 'Title'
        
        for i, finding in enumerate(findings):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = finding.get('severity', 'Medium')
            row_cells[2].text = finding.get('title', 'N/A')

        # Detailed Findings
        doc.add_heading('Detailed Findings', level=1)
        for i, finding in enumerate(findings):
            doc.add_heading(f"{i+1}. {finding['title']}", level=2)
            doc.add_paragraph(f"Severity: {finding['severity']}")
            doc.add_paragraph(f"Description: {finding['description']}")
            doc.add_paragraph(f"Remediation: {finding.get('remediation', 'N/A')}")

        # Save
        filename = f"Report_{client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        path = os.path.join(self.output_dir, filename)
        doc.save(path)
        return path
